import docx
import docx.shared

doc = docx.Document()

default_style = doc.styles['Normal']
default_style.font.name = 'Arial'
default_style.font.size = docx.shared.Pt(14)

header_style = doc.styles.add_style('MyHeader', docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
header_style.font.name = 'Arial'
header_style.font.size = docx.shared.Pt(30)
header_style.font.bold = True
header_style.font.color.rgb = docx.shared.RGBColor(150, 0, 0)
header_style.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

sub_header_style = doc.styles.add_style('MySubHeader', docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
sub_header_style.font.name = 'Arial'
sub_header_style.font.size = docx.shared.Pt(10)
sub_header_style.font.italic = True
sub_header_style.font.color.rgb = docx.shared.RGBColor(0, 0, 0)
sub_header_style.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

table_type_part_style = doc.styles.add_style('MyTableTypePart', docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
table_type_part_style.font.name = 'Arial'
table_type_part_style.font.size = docx.shared.Pt(14)
table_type_part_style.font.bold = True
table_type_part_style.font.color.rgb = docx.shared.RGBColor(50, 0, 0)
table_type_part_style.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

table_info_part_style = doc.styles.add_style('MyTableInfoPart', docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
table_info_part_style.font.name = 'Arial'
table_info_part_style.font.size = docx.shared.Pt(14)
table_info_part_style.font.color.rgb = docx.shared.RGBColor(0, 0, 0)
table_info_part_style.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

head = doc.add_paragraph("Tyler Durden", header_style)
sub_header = doc.add_paragraph('"Не разбив яиц, омлет не приготовишь."', sub_header_style)
sub_header.paragraph_format.space_after = docx.shared.Mm(10)
doc.add_picture("./pictures/tyler.jpg", width=docx.shared.Cm(15))

head = doc.add_paragraph("Информация", header_style)
head.paragraph_format.space_after = docx.shared.Mm(10)
info = [["Возраст", "30 лет"],
        ["Статус", "Жив"],
        ["Род занятий", "Координатор по возврату"],
        ["Жена", "Марла Сингер"],
        ["Дети", "Джуниор"],
        ["Прозвища", "Тайлер Дёрден, Корнелиус, Руперт, Ленни, Трэвис, Мистер Тейлор, Оззи, Бальтазар"]]
table = doc.add_table(len(info), 2, 'Table Grid')
table.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
for row in table.rows:
    row.height = docx.shared.Cm(2)
for i in range(len(info)):
    row = table.row_cells(i)
    row[0].paragraphs[0].add_run(info[i][0])
    row[0].paragraphs[0].style = table_info_part_style
    if i == len(info) - 1:
        names = [i.strip() for i in info[i][1].split(",")]
        column_count = 3
        rows_count = len(names) // column_count + int(bool(len(names) % column_count))
        sub_table = row[1].add_table(rows_count, column_count)
        sub_table.style = 'Table Grid'
        for r in sub_table.rows:
            r.height = docx.shared.Cm(2)
        for i in range(len(names)):
            sub_table.cell(i // column_count, i % column_count).paragraphs[0].add_run(names[i])
            sub_table.cell(i // column_count, i % column_count).paragraphs[0].style = table_info_part_style
        row[1].add_paragraph()
    else:
        row[1].paragraphs[0].add_run(info[i][1])
        row[1].paragraphs[0].style = table_info_part_style


doc.save("./tyler-durden.docx")